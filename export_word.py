# -*- coding: utf-8 -*-
"""
Genera el documento Word completo del proyecto:
  Análisis y Predicción del Caudal — Estación Pueblo Nuevo (DHIME-IDEAM)
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Aplica color de fondo a una celda de tabla."""
    shading = cell._element.get_or_add_tcPr()
    sh = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(sh)


def add_hyperlink(paragraph, url, text):
    """Agrega un hipervínculo al párrafo."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = paragraph._element.makeelement(qn('w:hyperlink'), {qn('r:id'): r_id})
    new_run = paragraph._element.makeelement(qn('w:r'), {})
    rPr = paragraph._element.makeelement(qn('w:rPr'), {})
    c = paragraph._element.makeelement(qn('w:color'), {qn('w:val'): '0563C1'})
    u = paragraph._element.makeelement(qn('w:u'), {qn('w:val'): 'single'})
    sz = paragraph._element.makeelement(qn('w:sz'), {qn('w:val'): '22'})
    rPr.append(c)
    rPr.append(u)
    rPr.append(sz)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)


def styled_paragraph(doc, text, style='Normal', bold=False, italic=False,
                     size=11, color=None, alignment=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(11)
        r.font.name = 'Calibri'
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
        r2.font.name = 'Calibri'
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.name = 'Calibri'
    return p


# ── DOCUMENTO ────────────────────────────────────────────────────────────────

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for s in doc.sections:
    s.top_margin = Cm(2.54)
    s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(2.54)
    s.right_margin = Cm(2.54)

# ═══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
styled_paragraph(doc, 'PROYECTO DE INVESTIGACIÓN', bold=True, size=14,
                 alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24,
                 color=(0, 51, 102))

styled_paragraph(doc, 'Análisis y Predicción del Caudal Medio Diario\nmediante Modelos de Series Temporales',
                 bold=True, size=18, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                 space_after=12, color=(0, 51, 102))

styled_paragraph(doc, 'Estación Pueblo Nuevo — DHIME / IDEAM (Colombia)\nCódigo 21117100 | Período 2010–2017',
                 size=13, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                 space_after=30, color=(80, 80, 80))

styled_paragraph(doc, 'Modelos evaluados: SARIMA · Holt-Winters · Prophet',
                 bold=True, size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                 space_after=6, color=(0, 102, 153))
styled_paragraph(doc, 'Variable: Caudal medio diario (m³/s) | Horizonte: 30 días',
                 size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                 space_after=48, color=(100, 100, 100))

styled_paragraph(doc, 'Fundamentos de Datos con Python',
                 bold=True, size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
styled_paragraph(doc, '2026', size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# TABLA DE CONTENIDO (manual)
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('Tabla de Contenido', level=1)
toc_items = [
    '1. Preguntas de Investigación',
    '2. Problema de Investigación',
    '3. Justificación',
    '4. Objetivo General',
    '5. Objetivos Específicos',
    '6. Marco Metodológico (Resumen por Semanas)',
    '   6.1 Semana 1 — Carga y Exploración',
    '   6.2 Semana 2 — Tratamiento de Datos Faltantes e Imputación',
    '   6.3 Semana 3 — EDA Avanzado y Estacionariedad',
    '   6.4 Semana 4 — Modelado Predictivo',
    '7. Resultados y Hallazgos Principales',
    '8. Conclusiones del Proyecto',
    '9. Trabajo Futuro',
    '10. Referencias Bibliográficas',
]
for item in toc_items:
    styled_paragraph(doc, item, size=11, space_after=3)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. PREGUNTAS DE INVESTIGACIÓN
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('1. Preguntas de Investigación', level=1)

doc.add_heading('Pregunta principal', level=2)
p = doc.add_paragraph()
r = p.add_run('¿Es posible pronosticar el caudal medio diario de la estación Pueblo Nuevo con un horizonte de 30 días, '
              'utilizando modelos de series temporales (SARIMA, Holt-Winters, Prophet) y aprovechando la estacionalidad anual identificada?')
r.italic = True
r.font.size = Pt(11)
r.font.name = 'Calibri'

doc.add_heading('Sub-preguntas', level=2)
sub_preguntas = [
    '¿Qué tan viable es el pronóstico del caudal medio diario a corto plazo (30 días) en estaciones hidrológicas del sistema DHIME-IDEAM mediante modelos estadísticos de series temporales?',
    '¿En qué medida la estacionalidad anual bimodal del régimen hidrológico colombiano contribuye a la precisión de modelos predictivos univariados (SARIMA, Holt-Winters, Prophet)?',
    '¿Cómo incide la calidad de los datos (gaps, outliers) y las técnicas de imputación aplicadas sobre la confiabilidad de los pronósticos de caudal?',
]
for sq in sub_preguntas:
    add_bullet(doc, sq)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 2. PROBLEMA DE INVESTIGACIÓN
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('2. Problema de Investigación', level=1)

styled_paragraph(doc,
    'La gestión de recursos hídricos en Colombia requiere pronósticos confiables de caudal para la toma de decisiones '
    'en alerta temprana por crecientes, planificación de riego y abastecimiento. Sin embargo, las estaciones hidrológicas '
    'del sistema DHIME-IDEAM presentan registros con vacíos temporales intermitentes (~9 % de datos faltantes), valores '
    'extremos asociados a eventos de crecida y una fuerte estacionalidad anual de régimen bimodal.',
    space_after=10)

styled_paragraph(doc,
    'A pesar de la disponibilidad de datos abiertos, no existe una evaluación comparativa sistemática de modelos clásicos '
    'de series temporales que aproveche estos patrones estacionales para generar pronósticos operacionales a corto plazo '
    'del caudal medio diario en estaciones de la zona andina colombiana.',
    space_after=10)

styled_paragraph(doc,
    'Este vacío impide que las autoridades ambientales y los gestores del recurso hídrico cuenten con herramientas '
    'cuantitativas reproducibles que traduzcan los datos históricos del IDEAM en pronósticos accionables para la '
    'planificación hidrológica regional.',
    space_after=10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 3. JUSTIFICACIÓN
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('3. Justificación', level=1)

justificaciones = [
    ('Relevancia hidrológica: ',
     'El caudal medio diario es una variable crítica para la gestión del recurso hídrico, la prevención de inundaciones '
     'y la planificación agrícola en Colombia. Disponer de pronósticos a 30 días permite anticipar condiciones de déficit '
     'o exceso hídrico.'),
    ('Disponibilidad de datos: ',
     'El sistema DHIME del IDEAM ofrece datos abiertos de miles de estaciones hidrometeorológicas. Sin embargo, su '
     'aprovechamiento analítico es limitado por la presencia de gaps temporales y la falta de pipelines reproducibles '
     'de análisis y pronóstico.'),
    ('Aporte metodológico: ',
     'El proyecto establece un pipeline completo y replicable — desde la carga cruda hasta el pronóstico — aplicable '
     'a cualquier estación del sistema DHIME, contribuyendo a la democratización del análisis hidrológico con herramientas '
     'de código abierto (Python, Pandas, Plotly, statsmodels, Prophet).'),
    ('Necesidad práctica: ',
     'Pronosticar con 30 días de anticipación permite la toma de decisiones informada en contextos de variabilidad '
     'climática (El Niño / La Niña), donde los patrones interanuales alteran significativamente el régimen de caudal.'),
    ('Contribución académica: ',
     'La comparación entre SARIMA, Holt-Winters y Prophet aplicados a datos reales colombianos genera conocimiento '
     'transferible sobre las fortalezas y limitaciones de cada modelo en contextos tropicales con estacionalidad bimodal.'),
]
for prefix, text in justificaciones:
    add_bullet(doc, text, bold_prefix=prefix)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 4. OBJETIVO GENERAL
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('4. Objetivo General', level=1)

styled_paragraph(doc,
    'Evaluar la capacidad predictiva de tres modelos de series temporales (SARIMA, Holt-Winters y Prophet) para '
    'pronosticar el caudal medio diario a un horizonte de 30 días en la estación Pueblo Nuevo (IDEAM, código 21117100), '
    'a partir de una serie limpia e imputada del período 2010–2017.',
    space_after=10)

# ═══════════════════════════════════════════════════════════════════════════════
# 5. OBJETIVOS ESPECÍFICOS
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('5. Objetivos Específicos', level=1)

objetivos = [
    ('OE1 — Preprocesamiento: ',
     'Construir un pipeline de preprocesamiento de la serie de caudal que incluya reindexación a frecuencia diaria, '
     'imputación de datos faltantes mediante interpolación lineal, detección y tratamiento de outliers por capping '
     '(P1–P99) y transformación logarítmica para estabilización de varianza.'),
    ('OE2 — Caracterización temporal: ',
     'Caracterizar los patrones temporales de la serie mediante descomposición STL, análisis de autocorrelación '
     '(ACF/PACF), pruebas de estacionariedad (ADF/KPSS) e identificación de la estructura estacional bimodal y la '
     'variabilidad interanual del caudal.'),
    ('OE3 — Modelado y comparación: ',
     'Entrenar, comparar y seleccionar el modelo con mejor desempeño predictivo entre SARIMA(1,1,1)(1,1,1)₃₀, '
     'Holt-Winters aditivo y Prophet con estacionalidad anual, evaluados sobre un período de prueba de 30 días '
     'mediante las métricas RMSE, MAE y MAPE.'),
]
for prefix, text in objetivos:
    add_bullet(doc, text, bold_prefix=prefix)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 6. MARCO METODOLÓGICO
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('6. Marco Metodológico (Resumen por Semanas)', level=1)

styled_paragraph(doc,
    'El proyecto sigue un pipeline secuencial de 4 fases (semanas), donde cada etapa produce los insumos requeridos '
    'por la siguiente. A continuación se describe cada fase con sus técnicas, herramientas y resultados clave.',
    space_after=12)

# ── Semana 1 ─────────────────────────────────────────────────────────────────
doc.add_heading('6.1 Semana 1 — Carga de Datos, Exploración Inicial y Visualización', level=2)

styled_paragraph(doc, 'Notebook: 01_carga_exploracion_caudal.ipynb', italic=True,
                 size=10, color=(100, 100, 100), space_after=8)

styled_paragraph(doc, 'Objetivo:', bold=True, space_after=2)
styled_paragraph(doc,
    'Cargar los datos de caudal medio diario desde el archivo descargaDhime.csv (IDEAM), explorar su estructura, '
    'eliminar columnas redundantes, diagnosticar la completitud temporal y generar las primeras visualizaciones interactivas.',
    space_after=8)

styled_paragraph(doc, 'Actividades realizadas:', bold=True, space_after=2)
actividades_s1 = [
    'Carga del CSV crudo (2,653 registros) y exploración con .info(), .describe().',
    'Revisión de la columna NivelAprobacion y eliminación de 6 columnas con valores únicos/constantes.',
    'Conversión de la columna Fecha a DatetimeIndex y renombramiento de Valor → Caudal.',
    'Estadísticas descriptivas: media, mediana, desviación estándar, percentiles.',
    'Diagnóstico de completitud temporal: identificación de gaps (días faltantes) por año.',
    'Visualización de la serie temporal completa con Plotly (interactiva).',
    'Distribución del caudal: histograma y boxplot — asimetría positiva detectada.',
    'Caudal promedio mensual: patrón estacional coherente con régimen hidrológico colombiano.',
    'Heatmap Año × Mes: visualización de promedios mensuales por año.',
]
for a in actividades_s1:
    add_bullet(doc, a)

styled_paragraph(doc, 'Hallazgos clave:', bold=True, space_after=2)
hallazgos_s1 = [
    'Datos de la estación PUEBLO NUEVO (código 21117100), período 2010–2017.',
    'La serie tiene días faltantes distribuidos en varios años.',
    'El caudal presenta asimetría positiva con presencia de valores extremos (crecientes).',
    'Se observa un patrón estacional coherente con el régimen hidrológico colombiano.',
]
for h in hallazgos_s1:
    add_bullet(doc, h)

# ── Semana 2 ─────────────────────────────────────────────────────────────────
doc.add_heading('6.2 Semana 2 — Tratamiento de Datos Faltantes, Imputación y Outliers', level=2)

styled_paragraph(doc, 'Notebook: 02_tratamiento_nan_imputacion_caudal.ipynb', italic=True,
                 size=10, color=(100, 100, 100), space_after=8)

styled_paragraph(doc, 'Objetivo:', bold=True, space_after=2)
styled_paragraph(doc,
    'Convertir la serie con gaps en una serie temporal diaria completa y limpia, lista para EDA avanzado y modelado.',
    space_after=8)

styled_paragraph(doc, 'Actividades realizadas:', bold=True, space_after=2)
actividades_s2 = [
    'Reindexación a frecuencia diaria completa (asfreq("D")) — se revelan ~269 días faltantes (completitud ~90.8 %).',
    'Visualización del patrón de datos faltantes: no aleatorios, concentrados en ciertos meses/años (fallas de estación).',
    'Comparación de 4 métodos de imputación: forward fill, interpolación lineal, interpolación temporal y media móvil.',
    'Selección de interpolación lineal por preservar mejor las estadísticas originales y generar transiciones suaves.',
    'Detección de outliers mediante IQR y Z-score.',
    'Tratamiento de outliers con capping (percentiles P1–P99) en lugar de eliminación (picos son eventos legítimos).',
    'Transformación logarítmica log(1+x) para reducir asimetría y estabilizar varianza.',
    'Transformación Box-Cox y diferenciación de primer orden.',
    'Verificación final de la serie limpia.',
    'Exportación del dataset limpio: caudal_limpio_diario.csv.',
]
for a in actividades_s2:
    add_bullet(doc, a)

styled_paragraph(doc, 'Hallazgos clave:', bold=True, space_after=2)
hallazgos_s2 = [
    'Los datos faltantes no son aleatorios — se concentran en ciertos períodos, sugiriendo fallas en la estación de medición.',
    'La interpolación lineal fue el método más equilibrado entre fidelidad estadística y suavidad.',
    'Los valores extremos altos corresponden a eventos hidrológicos legítimos (crecientes), por lo que se aplicó capping en lugar de eliminación.',
    'La transformación log(1+x) reduce significativamente la asimetría positiva.',
]
for h in hallazgos_s2:
    add_bullet(doc, h)

# ── Semana 3 ─────────────────────────────────────────────────────────────────
doc.add_heading('6.3 Semana 3 — EDA Avanzado, Estacionariedad y Pregunta de Investigación', level=2)

styled_paragraph(doc, 'Notebook: 03_eda_avanzado_estacionariedad_caudal.ipynb', italic=True,
                 size=10, color=(100, 100, 100), space_after=8)

styled_paragraph(doc, 'Objetivo:', bold=True, space_after=2)
styled_paragraph(doc,
    'Realizar un análisis exploratorio profundo de la serie limpia para identificar patrones temporales, evaluar '
    'estacionariedad y formular la pregunta de investigación que guiará el modelado predictivo.',
    space_after=8)

styled_paragraph(doc, 'Actividades realizadas:', bold=True, space_after=2)
actividades_s3 = [
    'Descomposición estacional STL: separación en Tendencia + Estacionalidad + Residuos.',
    'Cálculo de fuerza de tendencia y fuerza de estacionalidad.',
    'Análisis de autocorrelación (ACF): decaimiento lento, picos significativos en lag 365 (ciclo anual).',
    'Análisis de autocorrelación parcial (PACF): componente autorregresivo fuerte en lags 1-3.',
    'Prueba de Dickey-Fuller Aumentada (ADF) — serie original y diferenciada.',
    'Prueba KPSS — confirma no estacionariedad de la serie original.',
    'Formulación de la pregunta de investigación.',
    'Análisis de estacionalidad mensual y trimestral.',
    'Patrones interanuales y detección de variabilidad asociada a El Niño / La Niña.',
    'Distribución por épocas hidrológicas: seca vs lluviosa.',
    'Dashboard de resumen: 4 paneles (serie+tendencia, estacionalidad, ACF, promedios mensuales).',
]
for a in actividades_s3:
    add_bullet(doc, a)

styled_paragraph(doc, 'Hallazgos clave:', bold=True, space_after=2)
hallazgos_s3 = [
    'La serie presenta estacionalidad anual fuerte con régimen bimodal (dos picos), coherente con el patrón de precipitación andino colombiano.',
    'La ACF confirma ciclo anual (pico en lag 365); la PACF sugiere componente AR de orden bajo.',
    'Serie original no estacionaria (KPSS); con una diferenciación (d=1) se alcanza estacionariedad (ADF).',
    'Algunos años presentan caudales significativamente diferentes de la media global (variabilidad interanual).',
    'La clasificación seca/lluviosa muestra distribuciones claramente diferenciadas.',
]
for h in hallazgos_s3:
    add_bullet(doc, h)

# ── Semana 4 ─────────────────────────────────────────────────────────────────
doc.add_heading('6.4 Semana 4 — Modelado Predictivo (SARIMA, Holt-Winters, Prophet)', level=2)

styled_paragraph(doc, 'Notebook: 04_forecasting_sarima_hw_prophet_caudal.ipynb', italic=True,
                 size=10, color=(100, 100, 100), space_after=8)

styled_paragraph(doc, 'Objetivo:', bold=True, space_after=2)
styled_paragraph(doc,
    'Entrenar y comparar tres modelos de pronóstico de series temporales para predecir el caudal medio diario a 30 días, '
    'y responder la pregunta de investigación formulada en la Semana 3.',
    space_after=8)

styled_paragraph(doc, 'Actividades realizadas:', bold=True, space_after=2)
actividades_s4 = [
    'División Train/Test: últimos 30 días como conjunto de prueba.',
    'Modelo 1 — SARIMA(1,1,1)(1,1,1)₃₀: captura tendencia, autocorrelación y estacionalidad mensual aproximada.',
    'Modelo 2 — Holt-Winters (Suavización Exponencial Triple): descompone en nivel, tendencia y estacionalidad aditiva.',
    'Modelo 3 — Prophet (Meta): detecta automáticamente puntos de cambio y captura estacionalidad anual nativa.',
    'Comparación visual de los tres pronósticos vs valores reales en el período de test.',
    'Tabla comparativa de métricas: RMSE, MAE y MAPE para cada modelo.',
    'Selección del mejor modelo según menor RMSE.',
    'Reentrenamiento del modelo ganador con la serie completa.',
    'Pronóstico a futuro: 30 días más allá del último dato disponible.',
    'Respuesta fundamentada a la pregunta de investigación.',
]
for a in actividades_s4:
    add_bullet(doc, a)

styled_paragraph(doc, 'Descripción de los modelos:', bold=True, space_after=2)

add_bullet(doc,
    'Captura la tendencia y la estacionalidad mensual aproximada mediante SARIMAX. Su fortaleza es modelar '
    'autocorrelaciones a múltiples rezagos. Al usar un período estacional de 30 días (en lugar de 365, computacionalmente '
    'inviable), pierde parte de la componente anual.',
    bold_prefix='SARIMA(1,1,1)(1,1,1)₃₀: ')

add_bullet(doc,
    'Descompone la serie en nivel, tendencia y estacionalidad aditiva. Es eficiente computacionalmente y robusto '
    'para series con patrones estacionales regulares.',
    bold_prefix='Holt-Winters: ')

add_bullet(doc,
    'Diseñado para series con estacionalidad fuerte y datos faltantes. Detecta automáticamente puntos de cambio '
    'en la tendencia y captura la estacionalidad anual de forma nativa, lo que lo hace particularmente adecuado para '
    'datos hidrológicos.',
    bold_prefix='Prophet: ')

styled_paragraph(doc, 'Métricas de evaluación:', bold=True, space_after=2)

# Tabla de métricas (descripción)
table_m = doc.add_table(rows=4, cols=2)
table_m.alignment = WD_TABLE_ALIGNMENT.CENTER
table_m.style = 'Light Grid Accent 1'
headers_m = ['Métrica', 'Descripción']
for i, h in enumerate(headers_m):
    cell = table_m.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
metricas_desc = [
    ('RMSE (m³/s)', 'Raíz del error cuadrático medio — penaliza errores grandes'),
    ('MAE (m³/s)', 'Error absoluto medio — interpretación directa en m³/s'),
    ('MAPE (%)', 'Error porcentual absoluto medio — independiente de escala'),
]
for row_idx, (met, desc) in enumerate(metricas_desc, 1):
    table_m.rows[row_idx].cells[0].text = met
    table_m.rows[row_idx].cells[1].text = desc

doc.add_paragraph()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 7. RESULTADOS Y HALLAZGOS PRINCIPALES
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('7. Resultados y Hallazgos Principales', level=1)

# Tabla resumen del pipeline
doc.add_heading('Resumen del pipeline completo', level=2)

table_r = doc.add_table(rows=5, cols=4)
table_r.alignment = WD_TABLE_ALIGNMENT.CENTER
table_r.style = 'Light Grid Accent 1'
headers_r = ['Semana', 'Notebook', 'Objetivo', 'Resultado Clave']
for i, h in enumerate(headers_r):
    cell = table_r.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

pipeline_data = [
    ('1', '01_carga_exploracion_caudal', 'Carga, exploración y diagnóstico inicial',
     'Serie de ~2,653 registros (2010-2017), con gaps identificados'),
    ('2', '02_tratamiento_nan_imputacion_caudal', 'Limpieza, imputación y transformación',
     'Serie diaria completa, interpolación lineal, capping P1-P99'),
    ('3', '03_eda_avanzado_estacionariedad_caudal', 'EDA avanzado, descomposición y estacionariedad',
     'Estacionalidad anual confirmada (STL), serie no estacionaria, régimen bimodal'),
    ('4', '04_forecasting_sarima_hw_prophet_caudal', 'Pronóstico con SARIMA, Holt-Winters y Prophet',
     'Modelo óptimo identificado, pronóstico a 30 días generado'),
]
for row_idx, (sem, nb, obj, res) in enumerate(pipeline_data, 1):
    table_r.rows[row_idx].cells[0].text = sem
    table_r.rows[row_idx].cells[1].text = nb
    table_r.rows[row_idx].cells[2].text = obj
    table_r.rows[row_idx].cells[3].text = res

doc.add_paragraph()

styled_paragraph(doc, 'Hallazgos principales:', bold=True, space_after=4)

hallazgos_principales = [
    ('Calidad de datos: ', 'La estación PUEBLO NUEVO presenta datos con gaps intermitentes (~9 %) que fueron '
     'exitosamente imputados mediante interpolación lineal, preservando la estructura temporal.'),
    ('Patrón hidrológico: ', 'El caudal exhibe una marcada estacionalidad anual con régimen bimodal (dos picos), '
     'coherente con el patrón de precipitación de la zona andina colombiana.'),
    ('Capacidad predictiva: ', 'Los tres modelos evaluados logran capturar la tendencia general del caudal. El modelo '
     'con menor RMSE fue seleccionado como más adecuado para pronósticos operacionales.'),
    ('Valor práctico: ', 'El pipeline desarrollado es replicable para cualquier estación del sistema DHIME/IDEAM, '
     'contribuyendo a la gestión de recursos hídricos en Colombia.'),
]
for prefix, text in hallazgos_principales:
    add_bullet(doc, text, bold_prefix=prefix)

styled_paragraph(doc, '', space_after=4)
styled_paragraph(doc, 'Limitaciones identificadas:', bold=True, space_after=4)
limitaciones = [
    'Los eventos extremos (crecidas súbitas) son difíciles de predecir con modelos univariados.',
    'La estacionalidad anual dominante requiere períodos largos que aumentan la complejidad computacional (especialmente SARIMA).',
    'Variables exógenas (precipitación, temperatura, uso del suelo) podrían mejorar significativamente los pronósticos.',
]
for l in limitaciones:
    add_bullet(doc, l)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 8. CONCLUSIONES
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('8. Conclusiones del Proyecto', level=1)

conclusiones = [
    'Se demostró que es posible pronosticar el caudal medio diario de la estación Pueblo Nuevo a un horizonte de '
    '30 días utilizando modelos clásicos de series temporales, con diferentes grados de precisión según el modelo.',

    'El pipeline completo — desde carga cruda hasta pronóstico — es reproducible y aplicable a cualquier estación '
    'del sistema DHIME-IDEAM colombiano, lo cual lo convierte en una herramienta de utilidad práctica para la '
    'gestión hídrica.',

    'La interpolación lineal demostró ser el método más adecuado para imputar los ~269 días faltantes, preservando '
    'las estadísticas originales de la serie sin introducir artefactos significativos.',

    'La descomposición STL confirmó una estacionalidad anual fuerte y un régimen bimodal, componentes que son '
    'capturados de manera diferente por cada modelo: Prophet los maneja de forma nativa, Holt-Winters mediante '
    'componentes aditivos y SARIMA a través de diferenciación estacional.',

    'La comparación de métricas (RMSE, MAE, MAPE) permite una selección objetiva del modelo óptimo, '
    'evitando decisiones subjetivas en la elección de herramientas de pronóstico.',

    'El tratamiento de outliers mediante capping (P1–P99) en lugar de eliminación fue una decisión clave, ya que '
    'los picos de caudal representan eventos hidrológicos reales que deben ser preservados para el modelado.',
]
for i, c in enumerate(conclusiones, 1):
    add_bullet(doc, c, bold_prefix="{}. ".format(i))

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 9. TRABAJO FUTURO
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('9. Trabajo Futuro', level=1)

trabajo_futuro = [
    ('Variables exógenas: ', 'Incorporar datos de precipitación, temperatura y uso del suelo como regresores '
     'externos en modelos SARIMAX o Prophet con regresores.'),
    ('Aprendizaje profundo: ', 'Explorar modelos LSTM, GRU y Transformer para capturar patrones no lineales '
     'en la serie de caudal.'),
    ('Validación cruzada temporal: ', 'Implementar esquemas de validación con ventanas deslizantes para una '
     'evaluación más robusta del desempeño predictivo.'),
    ('Análisis multi-estación: ', 'Extender el análisis a múltiples estaciones del sistema DHIME para comparación '
     'regional y detección de patrones espaciales.'),
    ('Pronóstico ensemble: ', 'Combinar los pronósticos de los tres modelos mediante técnicas ensemble '
     '(promedio ponderado, stacking) para mejorar la precisión.'),
]
for prefix, text in trabajo_futuro:
    add_bullet(doc, text, bold_prefix=prefix)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 10. REFERENCIAS BIBLIOGRÁFICAS
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('10. Referencias Bibliográficas', level=1)

styled_paragraph(doc,
    'A continuación se listan 20 artículos científicos relevantes al proyecto, organizados por temática. '
    'Cada referencia incluye un enlace al artículo original.',
    space_after=12)

# ── Grupo 1: Pronóstico con SARIMA/HW/Prophet ──
doc.add_heading('Pronóstico de caudal con modelos de series de tiempo', level=2)

refs_grupo1 = [
    ('Short-Term Streamflow Forecasting for River Management',
     'https://www.mdpi.com/2306-5338/13/3/82', 'MDPI Hydrology'),
    ('Assessment of Time Series Models for Mean Discharge',
     'https://www.mdpi.com/2306-5338/10/11/208', 'MDPI Hydrology'),
    ('A hybrid SARIMA-Prophet model for predicting historical streamflow',
     'https://link.springer.com/article/10.1007/s42452-024-06083-x', 'Springer'),
    ('A data-driven approach to river discharge forecasting',
     'https://www.sciencedirect.com/science/article/pii/S2590123024002974', 'ScienceDirect'),
    ('Comparison between SARIMA and Holt-Winters for forecasting monthly streamflow (Cuba)',
     'https://link.springer.com/article/10.1007/s42452-021-04667-5', 'Springer'),
    ('Prophet-Based AI vs. Seasonal Auto-Regressive Models',
     'https://www.mdpi.com/2073-4441/17/24/3551', 'MDPI Water'),
]

refs_grupo2_title = 'Pronóstico hidrológico con métodos avanzados y comparaciones'
refs_grupo2 = [
    ('Evaluating Statistical and ML Models for River Flow Forecasting (Prophet, XGBoost, Random Forest)',
     'https://www.researchgate.net/publication/397317558', 'ResearchGate'),
    ('Hydrological time series forecasting using simple combinations: Big data testing',
     'https://www.sciencedirect.com/science/article/abs/pii/S002216942030665X', 'ScienceDirect'),
    ('Comparative assessment of ML models for daily streamflow prediction',
     'https://www.nature.com/articles/s41598-026-38969-8', 'Nature Scientific Reports'),
    ('Intercomparison of deep learning models in predicting streamflow patterns (CMIP6)',
     'https://www.nature.com/articles/s41598-024-63989-7', 'Nature Scientific Reports'),
    ('Comparative analysis of deep learning and ML for one-day-ahead streamflow forecasting',
     'https://www.sciencedirect.com/science/article/pii/S221458182500374X', 'ScienceDirect'),
    ('Effectiveness of three ML models for daily streamflow prediction and uncertainty',
     'https://www.sciencedirect.com/science/article/pii/S2589914724000860', 'ScienceDirect'),
    ('Exploring deep learning for streamflow forecasting: seasonal and perennial rivers',
     'https://www.sciencedirect.com/science/article/abs/pii/S0957417424010054', 'ScienceDirect'),
]

refs_grupo3_title = 'Holt-Winters en hidrología'
refs_grupo3 = [
    ('Flood Forecasting using Holt-Winters Exponential Smoothing and GIS',
     'https://www.researchgate.net/publication/322003786', 'ResearchGate'),
    ('The use of Holt-Winters method for forecasting the amount of water inflow',
     'https://www.researchgate.net/publication/303908593', 'ResearchGate'),
]

refs_grupo4_title = 'Imputación de datos faltantes en series hidrológicas'
refs_grupo4 = [
    ('Improving Linear Interpolation of Missing Hydrological Data by Applying Integrated AR Models',
     'https://link.springer.com/article/10.1007/s11269-023-03625-7', 'Springer'),
    ('Comparing single- and two-segment statistical models with a conceptual model',
     'https://www.sciencedirect.com/science/article/pii/S1364815216305138', 'ScienceDirect'),
]

refs_grupo5_title = 'Hidrología en Colombia (IDEAM / contexto regional)'
refs_grupo5 = [
    ('Seasonal streamflow prediction in Colombia using atmospheric and oceanic indices',
     'https://www.sciencedirect.com/science/article/abs/pii/S0022169416301846', 'ScienceDirect'),
    ('Spatio-temporal analysis of hydrological response to land cover changes — Chicú river, Colombia',
     'https://www.sciencedirect.com/science/article/pii/S2405844021014614', 'ScienceDirect'),
    ('Magdalena river: interannual variability (1975-1995) and revised water discharge',
     'https://www.sciencedirect.com/science/article/abs/pii/S0022169400002699', 'ScienceDirect'),
]

all_ref_groups = [
    (None, refs_grupo1),
    (refs_grupo2_title, refs_grupo2),
    (refs_grupo3_title, refs_grupo3),
    (refs_grupo4_title, refs_grupo4),
    (refs_grupo5_title, refs_grupo5),
]

ref_num = 1
for group_title, refs in all_ref_groups:
    if group_title:
        doc.add_heading(group_title, level=2)
    for title, url, source in refs:
        p = doc.add_paragraph()
        r = p.add_run("[{}] ".format(ref_num))
        r.bold = True
        r.font.size = Pt(11)
        r.font.name = 'Calibri'
        r2 = p.add_run(title + ". ")
        r2.font.size = Pt(11)
        r2.font.name = 'Calibri'
        add_hyperlink(p, url, url)
        p.paragraph_format.space_after = Pt(6)
        ref_num += 1

# ═══════════════════════════════════════════════════════════════════════════════
# GUARDAR
# ═══════════════════════════════════════════════════════════════════════════════

output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                           'Proyecto_Caudal_Pueblo_Nuevo_Completo.docx')
doc.save(output_path)
print("Documento generado exitosamente: {}".format(output_path))
